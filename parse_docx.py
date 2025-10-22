#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
docx文件解析脚本，用于将docx文档转换为HTML格式
"""

import os
import sys
import json
from docx import Document
from datetime import datetime

# 文章元数据映射
ARTICLE_METADATA = {
    "12.14雪夜.docx": {
        "id": "1",
        "title": "12.14雪夜",
        "date": "2023-12-14",
        "category": "文学创作",
        "description": "一场雪夜的思考与感悟"
    },
    "mindcopy标准介绍及其示例.docx": {
        "id": "2",
        "title": "mindcopy标准介绍及其示例",
        "date": "2023-11-20",
        "category": "技术标准",
        "description": "关于意识复制技术的标准与实践"
    },
    "宇宙的意识.docx": {
        "id": "3",
        "title": "宇宙的意识",
        "date": "2023-10-05",
        "category": "哲学思考",
        "description": "探索宇宙与意识的关系"
    },
    "青春诗 2.0.docx": {
        "id": "4",
        "title": "青春诗 2.0",
        "date": "2023-09-15",
        "category": "诗歌创作",
        "description": "关于青春的诗意表达"
    }
}

def clean_text(text):
    """清理文本，去除多余的空白字符"""
    return ' '.join(text.split()) if text else ''

def docx_to_html(file_path):
    """将docx文件转换为HTML格式"""
    try:
        doc = Document(file_path)
        html_content = []
        
        for paragraph in doc.paragraphs:
            text = clean_text(paragraph.text)
            if not text:
                continue
                
            # 根据段落样式添加HTML标签
            if paragraph.style.name.startswith('Heading'):
                level = paragraph.style.name.replace('Heading ', '')
                if level.isdigit() and 1 <= int(level) <= 6:
                    html_content.append(f'<h{level}>{text}</h{level}>')
                else:
                    html_content.append(f'<p>{text}</p>')
            else:
                html_content.append(f'<p>{text}</p>')
        
        # 处理表格（简单实现）
        for table in doc.tables:
            html_content.append('<table>')
            for row in table.rows:
                html_content.append('<tr>')
                for cell in row.cells:
                    cell_text = clean_text(cell.text)
                    html_content.append(f'<td>{cell_text}</td>')
                html_content.append('</tr>')
            html_content.append('</table>')
        
        return '\n'.join(html_content)
    except Exception as e:
        print(f"解析文件 {file_path} 时出错: {str(e)}")
        return f'<p>解析错误: {str(e)}</p>'

def generate_articles_data(articles_dir):
    """生成所有文章的数据"""
    articles_data = []
    
    for filename in os.listdir(articles_dir):
        if filename.endswith('.docx'):
            file_path = os.path.join(articles_dir, filename)
            metadata = ARTICLE_METADATA.get(filename, {
                "id": str(len(articles_data) + 1),
                "title": filename.replace('.docx', ''),
                "date": datetime.now().strftime("%Y-%m-%d"),
                "category": "未分类",
                "description": "暂无描述"
            })
            
            # 解析文章内容
            content = docx_to_html(file_path)
            
            article = {
                **metadata,
                "filename": filename,
                "content": content
            }
            articles_data.append(article)
    
    # 按日期倒序排序
    articles_data.sort(key=lambda x: x['date'], reverse=True)
    
    return articles_data

def main():
    """主函数"""
    articles_dir = './aricle'
    if not os.path.exists(articles_dir):
        print(f"文章目录不存在: {articles_dir}")
        sys.exit(1)
    
    # 生成文章数据
    articles_data = generate_articles_data(articles_dir)
    
    # 输出为JSON格式
    output_file = 'articles_data.json'
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(articles_data, f, ensure_ascii=False, indent=2)
    
    print(f"文章数据已生成: {output_file}")
    print(f"共解析 {len(articles_data)} 篇文章")

if __name__ == "__main__":
    main()